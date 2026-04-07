from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_workflow_guide():
    doc = Document()
    
    # 标题
    title = doc.add_heading('“yuntian.huang的AI私人助理” Coze 对话流（Workflow）开发手册', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 第一章：核心设计思路
    doc.add_heading('1. 核心设计思路：意图分流架构', level=1)
    doc.add_paragraph('为了实现精准回复并提升用户体验，我们将采用 Coze 的“工作流（Workflow）”模式。核心逻辑是通过一个“意图分类器”将用户问题导向三个特定的处理分支：')
    
    # 场景列表
    doc.add_paragraph('分支 A：打招呼寒暄（你好、你是谁、你有什么技能等）', style='List Bullet')
    doc.add_paragraph('分支 B：私人信息类（简历、教育背景、实习经历、技能等）', style='List Bullet')
    doc.add_paragraph('分支 C：其他类（无法处理的问题、泛泛而谈等）', style='List Bullet')

    # 逻辑流程图（文本模拟）
    doc.add_heading('2. 对话流逻辑架构图', level=1)
    logic_box = doc.add_table(rows=1, cols=1)
    logic_box.style = 'Table Grid'
    cell = logic_box.rows[0].cells[0]
    cell.text = """[Start 节点] --> [意图分类节点 (LLM)] --> [Condition 选择器]
                                         |
    +------------------------------------+------------------------------------+
    |                                    |                                    |
[分支 A: 寒暄回复]                   [分支 B: 知识库检索]                 [分支 C: 通用回复]
    |                                    |                                    |
    +------------------------------------+------------------------------------+
                                         |
                                   [End 节点]"""

    # 第二章：详细配置步骤
    doc.add_heading('3. 详细配置步骤', level=1)
    
    doc.add_heading('步骤一：创建 Workflow', level=2)
    doc.add_paragraph('1. 登录 Coze 控制台，选择“工作流（Workflow）”。', style='List Number')
    doc.add_paragraph('2. 点击“创建工作流”，命名为 yuntian_assistant_workflow。', style='List Number')

    doc.add_heading('步骤二：配置“意图分类”节点', level=2)
    doc.add_paragraph('在 Start 节点后添加一个 LLM 节点，用于分析用户意图。其 Prompt 配置如下：')
    
    prompt_table = doc.add_table(rows=1, cols=1)
    prompt_table.style = 'Table Grid'
    p_cell = prompt_table.rows[0].cells[0]
    p_cell.text = """# Role
你是一个专业意图分类专家。

# Task
将用户输入分为以下三类并仅输出类别字母：
- A: 打招呼寒暄（如：你好、自我介绍、询问机器人身份）。
- B: 个人履历咨询（如：在哪上学、有什么实习经历、九号公司做了什么、有哪些AI技能）。
- C: 其他话题（无法理解或不相关的通用问题）。

# Constraint
仅输出一个字母 A, B, 或 C。"""

    doc.add_heading('步骤三：配置各逻辑分支', level=2)
    
    # Branch A
    doc.add_paragraph('分支 A (寒暄)：', style='List Bullet')
    doc.add_paragraph('连接一个 LLM 节点。Prompt 建议：作为“黄运添的 AI 助理”，向用户展示专业且亲和的形象，告知用户你可以回答关于其求职相关的背景信息。')
    
    # Branch B
    doc.add_paragraph('分支 B (私人信息)：', style='List Bullet')
    doc.add_paragraph('1. 连接“知识库（Knowledge）”节点：搜索您的 PRD/简历库。')
    doc.add_paragraph('2. 连接 LLM 节点：基于搜索结果，以第一人称或专业助理角度总结输出（例如：黄运添在吉林农业大学就读期间...）。')
    
    # Branch C
    doc.add_paragraph('分支 C (其他)：', style='List Bullet')
    doc.add_paragraph('连接一个 LLM 节点。Prompt 建议：礼貌回应，提示用户您可以回答简历相关问题。')

    # 第三章：调试与发布
    doc.add_heading('4. 调试与发布', level=1)
    doc.add_paragraph('1. 试运行：点击 Workflow 顶部的“试运行”，输入“你在九号公司做了什么？”检查是否正确进入分支 B。', style='List Number')
    doc.add_paragraph('2. 发布：将该 Workflow 绑定至 Bot 的“人设对话键”或设置为触发指令。', style='List Number')

    # 保存文件
    file_name = 'yuntian.huang的AI私人助理-Coze开发流程-v2.docx'
    doc.save(file_name)
    print(f"File created: {file_name}")

if __name__ == "__main__":
    create_workflow_guide()
